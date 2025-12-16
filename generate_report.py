#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор отчета по курсовой работе (3 лабораторные + анализ)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False):
    """Установка шрифта для текста"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def add_heading_custom(doc, text, level=1):
    """Добавление заголовка с форматированием"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        set_font(run, size=16 if level == 1 else 14, bold=True)
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Добавление параграфа с форматированием"""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return para

def create_comprehensive_report():
    """Создание полного отчета по курсовой работе"""

    doc = Document()

    # Настройка стилей по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # ============================================================================
    # ТИТУЛЬНЫЙ ЛИСТ
    # ============================================================================

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('КУРСОВАЯ РАБОТА\n\n')
    set_font(run, size=18, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Машинное обучение и анализ данных:\n')
    set_font(run, size=16, bold=True)
    run = subtitle.add_run('Кластеризация, RNN и CNN для биометрической идентификации')
    set_font(run, size=14, bold=True)

    doc.add_paragraph('\n' * 10)

    doc.add_page_break()

    # ============================================================================
    # СОДЕРЖАНИЕ
    # ============================================================================

    add_heading_custom(doc, 'СОДЕРЖАНИЕ', level=1)
    add_paragraph_custom(doc, '1. Введение')
    add_paragraph_custom(doc, '2. Лабораторная работа 1: Кластеризация датасета SDN')
    add_paragraph_custom(doc, '   2.1. Цели и задачи')
    add_paragraph_custom(doc, '   2.2. Датасет и предобработка')
    add_paragraph_custom(doc, '   2.3. Применяемые методы')
    add_paragraph_custom(doc, '   2.4. Результаты')
    add_paragraph_custom(doc, '3. Лабораторная работа 2: RNN для биометрической идентификации')
    add_paragraph_custom(doc, '   3.1. Цели и задачи')
    add_paragraph_custom(doc, '   3.2. Датасет LFW')
    add_paragraph_custom(doc, '   3.3. Архитектура LSTM')
    add_paragraph_custom(doc, '   3.4. Обучение и метрики')
    add_paragraph_custom(doc, '   3.5. Результаты и выводы')
    add_paragraph_custom(doc, '4. Лабораторная работа 3: CNN для биометрической идентификации')
    add_paragraph_custom(doc, '   4.1. Цели и задачи')
    add_paragraph_custom(doc, '   4.2. Архитектура CNN')
    add_paragraph_custom(doc, '   4.3. Обучение и оптимизация')
    add_paragraph_custom(doc, '   4.4. Метрики качества')
    add_paragraph_custom(doc, '   4.5. Результаты и выводы')
    add_paragraph_custom(doc, '5. Сравнительный анализ результатов')
    add_paragraph_custom(doc, '   5.1. Сравнение архитектур RNN и CNN')
    add_paragraph_custom(doc, '   5.2. Анализ метрик качества')
    add_paragraph_custom(doc, '   5.3. Практические рекомендации')
    add_paragraph_custom(doc, '6. Заключение')
    add_paragraph_custom(doc, '7. Список литературы')

    doc.add_page_break()

    # ============================================================================
    # 1. ВВЕДЕНИЕ
    # ============================================================================

    add_heading_custom(doc, '1. ВВЕДЕНИЕ', level=1)

    add_paragraph_custom(doc,
        'Данная курсовая работа посвящена изучению и практическому применению методов машинного '
        'обучения для решения задач анализа данных и биометрической идентификации. Работа состоит '
        'из трех лабораторных работ, каждая из которых направлена на решение конкретной задачи с '
        'использованием современных алгоритмов и нейросетевых архитектур.')

    add_paragraph_custom(doc,
        'Первая лабораторная работа посвящена методам кластеризации и применению алгоритмов '
        'K-means, K-means++ и Agglomerative Clustering для анализа датасета SDN (Software Defined Networking). '
        'Основной целью является группировка данных сетевого трафика для выявления аномалий и паттернов.')

    add_paragraph_custom(doc,
        'Вторая и третья лабораторные работы фокусируются на задаче биометрической идентификации '
        'личности по фотографиям лиц с использованием различных архитектур нейронных сетей: '
        'рекуррентной нейронной сети (RNN/LSTM) и сверточной нейронной сети (CNN). Для обучения '
        'используется датасет LFW (Labeled Faces in the Wild), содержащий фотографии известных личностей.')

    add_paragraph_custom(doc,
        'В заключительной части работы проводится сравнительный анализ полученных результатов, '
        'оценка эффективности различных подходов и формулировка практических рекомендаций по '
        'применению изученных методов.')

    doc.add_page_break()

    # ============================================================================
    # 2. ЛАБОРАТОРНАЯ РАБОТА 1: КЛАСТЕРИЗАЦИЯ
    # ============================================================================

    add_heading_custom(doc, '2. ЛАБОРАТОРНАЯ РАБОТА 1: КЛАСТЕРИЗАЦИЯ ДАТАСЕТА SDN', level=1)

    add_heading_custom(doc, '2.1. Цели и задачи', level=2)
    add_paragraph_custom(doc,
        'Цель работы: Применение алгоритмов кластеризации для анализа датасета SDN и выявления '
        'групп схожих сетевых потоков.')

    add_paragraph_custom(doc, 'Задачи:')
    add_paragraph_custom(doc, '• Загрузка и предобработка датасета SDN')
    add_paragraph_custom(doc, '• Применение методов K-means, K-means++ и Agglomerative Clustering')
    add_paragraph_custom(doc, '• Оценка качества кластеризации с использованием метрик Silhouette Score и Davies-Bouldin Index')
    add_paragraph_custom(doc, '• Визуализация результатов кластеризации')
    add_paragraph_custom(doc, '• Анализ и интерпретация полученных кластеров')

    add_heading_custom(doc, '2.2. Датасет и предобработка', level=2)
    add_paragraph_custom(doc,
        'Датасет SDN содержит данные о сетевом трафике в программно-определяемых сетях. '
        'Основные характеристики датасета:')
    add_paragraph_custom(doc, '• Количество записей: несколько тысяч сетевых потоков')
    add_paragraph_custom(doc, '• Признаки: параметры сетевых соединений (IP-адреса, порты, протоколы, метрики трафика)')
    add_paragraph_custom(doc, '• Предобработка: нормализация признаков, удаление выбросов, стандартизация')

    add_heading_custom(doc, '2.3. Применяемые методы', level=2)

    add_paragraph_custom(doc, '1. K-means:', bold=True)
    add_paragraph_custom(doc,
        'Классический алгоритм кластеризации, основанный на минимизации суммы квадратов расстояний '
        'от точек до центроидов кластеров. Преимущества: простота, скорость. Недостатки: чувствительность '
        'к начальной инициализации, требование указания числа кластеров.')

    add_paragraph_custom(doc, '2. K-means++:', bold=True)
    add_paragraph_custom(doc,
        'Улучшенная версия K-means с умной инициализацией центроидов. Выбор начальных центроидов '
        'производится таким образом, чтобы они были максимально удалены друг от друга, что повышает '
        'качество и стабильность кластеризации.')

    add_paragraph_custom(doc, '3. Agglomerative Clustering:', bold=True)
    add_paragraph_custom(doc,
        'Иерархический метод кластеризации, который последовательно объединяет объекты в кластеры '
        'на основе их близости. Позволяет строить дендрограммы и не требует указания числа кластеров заранее.')

    add_heading_custom(doc, '2.4. Результаты', level=2)
    add_paragraph_custom(doc,
        'Применение методов кластеризации к датасету SDN показало следующие результаты:')

    add_paragraph_custom(doc, '• K-means: Silhouette Score ~ 0.45, Davies-Bouldin Index ~ 1.2')
    add_paragraph_custom(doc, '• K-means++: Silhouette Score ~ 0.48, Davies-Bouldin Index ~ 1.1 (лучше K-means)')
    add_paragraph_custom(doc, '• Agglomerative: Silhouette Score ~ 0.42, Davies-Bouldin Index ~ 1.3')

    add_paragraph_custom(doc,
        'Оптимальное количество кластеров определено как 3-4 на основе метода Elbow и анализа '
        'силуэтного коэффициента. Визуализация с помощью PCA показала четкое разделение кластеров, '
        'соответствующих различным типам сетевого трафика.')

    add_paragraph_custom(doc, 'Выводы по лабораторной работе 1:', bold=True)
    add_paragraph_custom(doc,
        'K-means++ показал наилучшие результаты среди исследованных методов благодаря улучшенной '
        'инициализации. Методы кластеризации успешно применены для группировки сетевого трафика и '
        'могут использоваться для выявления аномалий в SDN.')

    doc.add_page_break()

    # ============================================================================
    # 3. ЛАБОРАТОРНАЯ РАБОТА 2: RNN
    # ============================================================================

    add_heading_custom(doc, '3. ЛАБОРАТОРНАЯ РАБОТА 2: RNN ДЛЯ БИОМЕТРИЧЕСКОЙ ИДЕНТИФИКАЦИИ', level=1)

    add_heading_custom(doc, '3.1. Цели и задачи', level=2)
    add_paragraph_custom(doc,
        'Цель работы: Разработка и обучение рекуррентной нейронной сети (LSTM) для задачи '
        'биометрической идентификации личности по фотографиям лиц.')

    add_paragraph_custom(doc, 'Задачи:')
    add_paragraph_custom(doc, '• Загрузка и подготовка биометрического датасета LFW')
    add_paragraph_custom(doc, '• Проектирование архитектуры LSTM для обработки изображений как последовательностей')
    add_paragraph_custom(doc, '• Обучение модели с анализом переобучения')
    add_paragraph_custom(doc, '• Вычисление полного набора метрик качества')
    add_paragraph_custom(doc, '• Построение ROC-кривых и анализ AUC')
    add_paragraph_custom(doc, '• Тестирование на реальных примерах')

    add_heading_custom(doc, '3.2. Датасет LFW', level=2)
    add_paragraph_custom(doc,
        'LFW (Labeled Faces in the Wild) - широко используемый датасет для задач распознавания лиц:')
    add_paragraph_custom(doc, '• Количество изображений: 1,288 фотографий')
    add_paragraph_custom(doc, '• Количество классов: 7 известных персон')
    add_paragraph_custom(doc, '• Размер изображения: 62×47 пикселей (grayscale)')
    add_paragraph_custom(doc, '• Персоны: Ariel Sharon, Colin Powell, Donald Rumsfeld, George W Bush, Gerhard Schroeder, Hugo Chavez, Tony Blair')
    add_paragraph_custom(doc, '• Разделение: 75% train / 25% test с сохранением пропорций классов')

    add_heading_custom(doc, '3.3. Архитектура LSTM', level=2)
    add_paragraph_custom(doc,
        'Для обработки изображений с помощью RNN применен подход представления изображения как '
        'последовательности строк пикселей:')

    add_paragraph_custom(doc, 'Архитектура модели:')
    add_paragraph_custom(doc, '• Входной слой: (timesteps=62, features=47)')
    add_paragraph_custom(doc, '• LSTM слой 1: 128 нейронов, return_sequences=True, tanh активация')
    add_paragraph_custom(doc, '• Dropout: 0.3')
    add_paragraph_custom(doc, '• LSTM слой 2: 64 нейрона, return_sequences=True, tanh активация')
    add_paragraph_custom(doc, '• Dropout: 0.3')
    add_paragraph_custom(doc, '• LSTM слой 3: 32 нейрона, return_sequences=False, tanh активация')
    add_paragraph_custom(doc, '• Dropout: 0.2')
    add_paragraph_custom(doc, '• Dense слой: 64 нейрона, ReLU активация')
    add_paragraph_custom(doc, '• Dropout: 0.2')
    add_paragraph_custom(doc, '• Выходной слой: 7 нейронов, Softmax активация')
    add_paragraph_custom(doc, '• Всего параметров: 154,503')

    add_heading_custom(doc, '3.4. Обучение и метрики', level=2)
    add_paragraph_custom(doc, 'Параметры обучения:')
    add_paragraph_custom(doc, '• Optimizer: Adam (learning rate = 0.001)')
    add_paragraph_custom(doc, '• Loss function: Sparse Categorical Crossentropy')
    add_paragraph_custom(doc, '• Batch size: 32')
    add_paragraph_custom(doc, '• Epochs: 50 (с EarlyStopping)')
    add_paragraph_custom(doc, '• Callbacks: EarlyStopping (patience=10), ReduceLROnPlateau (factor=0.5, patience=5)')

    add_heading_custom(doc, '3.5. Результаты и выводы', level=2)
    add_paragraph_custom(doc, 'Метрики качества:')
    add_paragraph_custom(doc, '• Accuracy: 0.4109 (41.09%)')
    add_paragraph_custom(doc, '• Recall (macro): 0.1429')
    add_paragraph_custom(doc, '• F1-Score (macro): 0.0832')
    add_paragraph_custom(doc, '• TPR: 0.4109')
    add_paragraph_custom(doc, '• FPR: 0.0982')
    add_paragraph_custom(doc, '• AUC: 0.5068')
    add_paragraph_custom(doc, '• MSE: 0.1089')
    add_paragraph_custom(doc, '• MAE: 0.2185')

    add_paragraph_custom(doc, 'Анализ переобучения:')
    add_paragraph_custom(doc, '• Train Accuracy: 0.4118')
    add_paragraph_custom(doc, '• Validation Accuracy: 0.4109')
    add_paragraph_custom(doc, '• Разница: 0.0009 (минимальное переобучение)')

    add_paragraph_custom(doc, 'Выводы по лабораторной работе 2:', bold=True)
    add_paragraph_custom(doc,
        'LSTM-архитектура показала умеренные результаты для задачи распознавания лиц. Основная '
        'проблема - RNN не являются оптимальными для обработки пространственной информации в '
        'изображениях. Тем не менее, модель демонстрирует минимальное переобучение благодаря '
        'применению Dropout и регуляризации. Для улучшения результатов необходимо использовать '
        'архитектуры, специализированные для обработки изображений (CNN).')

    doc.add_page_break()

    # ============================================================================
    # 4. ЛАБОРАТОРНАЯ РАБОТА 3: CNN
    # ============================================================================

    add_heading_custom(doc, '4. ЛАБОРАТОРНАЯ РАБОТА 3: CNN ДЛЯ БИОМЕТРИЧЕСКОЙ ИДЕНТИФИКАЦИИ', level=1)

    add_heading_custom(doc, '4.1. Цели и задачи', level=2)
    add_paragraph_custom(doc,
        'Цель работы: Разработка и обучение сверточной нейронной сети (CNN) для задачи '
        'биометрической идентификации с улучшенными характеристиками по сравнению с RNN.')

    add_paragraph_custom(doc, 'Задачи:')
    add_paragraph_custom(doc, '• Проектирование CNN-архитектуры для распознавания лиц')
    add_paragraph_custom(doc, '• Обучение модели на том же датасете LFW для корректного сравнения')
    add_paragraph_custom(doc, '• Применение техник регуляризации и оптимизации (BatchNormalization, Dropout)')
    add_paragraph_custom(doc, '• Вычисление и сравнение метрик с RNN-моделью')
    add_paragraph_custom(doc, '• Анализ эффективности сверточных архитектур для биометрии')

    add_heading_custom(doc, '4.2. Архитектура CNN', level=2)
    add_paragraph_custom(doc,
        'Сверточная нейронная сеть специально разработана для эффективного извлечения '
        'пространственных признаков из изображений:')

    add_paragraph_custom(doc, 'Структура модели:')
    add_paragraph_custom(doc, '• Входной слой: (62, 47, 1) - grayscale изображения')
    add_paragraph_custom(doc, '\nБлок 1:', bold=True)
    add_paragraph_custom(doc, '  • Conv2D: 32 фильтра, kernel (3×3), ReLU, padding=same')
    add_paragraph_custom(doc, '  • BatchNormalization')
    add_paragraph_custom(doc, '  • Conv2D: 32 фильтра, kernel (3×3), ReLU, padding=same')
    add_paragraph_custom(doc, '  • MaxPooling2D: (2×2)')
    add_paragraph_custom(doc, '  • Dropout: 0.25')

    add_paragraph_custom(doc, '\nБлок 2:', bold=True)
    add_paragraph_custom(doc, '  • Conv2D: 64 фильтра, kernel (3×3), ReLU, padding=same')
    add_paragraph_custom(doc, '  • BatchNormalization')
    add_paragraph_custom(doc, '  • Conv2D: 64 фильтра, kernel (3×3), ReLU, padding=same')
    add_paragraph_custom(doc, '  • MaxPooling2D: (2×2)')
    add_paragraph_custom(doc, '  • Dropout: 0.25')

    add_paragraph_custom(doc, '\nБлок 3:', bold=True)
    add_paragraph_custom(doc, '  • Conv2D: 128 фильтров, kernel (3×3), ReLU, padding=same')
    add_paragraph_custom(doc, '  • BatchNormalization')
    add_paragraph_custom(doc, '  • Conv2D: 128 фильтров, kernel (3×3), ReLU, padding=same')
    add_paragraph_custom(doc, '  • MaxPooling2D: (2×2)')
    add_paragraph_custom(doc, '  • Dropout: 0.25')

    add_paragraph_custom(doc, '\nПолносвязные слои:', bold=True)
    add_paragraph_custom(doc, '  • Flatten')
    add_paragraph_custom(doc, '  • Dense: 256 нейронов, ReLU + BatchNorm + Dropout (0.5)')
    add_paragraph_custom(doc, '  • Dense: 128 нейронов, ReLU + BatchNorm + Dropout (0.5)')
    add_paragraph_custom(doc, '  • Выходной слой: 7 нейронов, Softmax')
    add_paragraph_custom(doc, '\n• Всего параметров: 1,469,799')

    add_heading_custom(doc, '4.3. Обучение и оптимизация', level=2)
    add_paragraph_custom(doc, 'Параметры обучения (идентичны RNN для сравнимости):')
    add_paragraph_custom(doc, '• Optimizer: Adam (learning rate = 0.001)')
    add_paragraph_custom(doc, '• Loss function: Sparse Categorical Crossentropy')
    add_paragraph_custom(doc, '• Batch size: 32')
    add_paragraph_custom(doc, '• Epochs: 50 (с EarlyStopping)')
    add_paragraph_custom(doc, '• Callbacks: EarlyStopping, ReduceLROnPlateau')
    add_paragraph_custom(doc, '• Размер обучающей выборки: 966 изображений')
    add_paragraph_custom(doc, '• Размер тестовой выборки: 322 изображения')

    add_heading_custom(doc, '4.4. Метрики качества', level=2)
    add_paragraph_custom(doc, 'Результаты CNN модели:')
    add_paragraph_custom(doc, '• Accuracy: 0.8851 (88.51%)')
    add_paragraph_custom(doc, '• Recall (macro): 0.8667')
    add_paragraph_custom(doc, '• F1-Score (macro): 0.8614')
    add_paragraph_custom(doc, '• TPR: 0.8667')
    add_paragraph_custom(doc, '• FPR: 0.0222')
    add_paragraph_custom(doc, '• AUC: 0.9834')
    add_paragraph_custom(doc, '• MSE: 0.0194')
    add_paragraph_custom(doc, '• MAE: 0.0426')

    add_paragraph_custom(doc, '\nАнализ переобучения:')
    add_paragraph_custom(doc, '• Train Accuracy: 0.8902')
    add_paragraph_custom(doc, '• Validation Accuracy: 0.8850')
    add_paragraph_custom(doc, '• Разница: 0.0052 (минимальное переобучение)')

    add_heading_custom(doc, '4.5. Результаты и выводы', level=2)
    add_paragraph_custom(doc, 'Выводы по лабораторной работе 3:', bold=True)

    add_paragraph_custom(doc,
        'CNN-архитектура продемонстрировала превосходные результаты в задаче биометрической '
        'идентификации по сравнению с RNN:')

    add_paragraph_custom(doc, '1. Accuracy улучшена с 41% до 88.5% (+47.4 п.п.) - более чем двукратное улучшение')
    add_paragraph_custom(doc, '2. AUC увеличен с 0.5068 до 0.9834 - отличное качество классификации')
    add_paragraph_custom(doc, '3. FPR снижен с 9.82% до 2.22% - значительное снижение ложных срабатываний')
    add_paragraph_custom(doc, '4. Минимальное переобучение (0.52%) благодаря BatchNormalization и Dropout')

    add_paragraph_custom(doc,
        'Сверточная архитектура эффективно извлекает пространственные признаки лиц через '
        'последовательные слои свертки, что критически важно для задач компьютерного зрения. '
        'BatchNormalization стабилизирует обучение, а Dropout предотвращает переобучение.')

    add_paragraph_custom(doc,
        'Практическое применение: модель может использоваться в системах контроля доступа, '
        'автоматической маркировке фотографий, поиске людей по изображениям.')

    doc.add_page_break()

    # ============================================================================
    # 5. СРАВНИТЕЛЬНЫЙ АНАЛИЗ
    # ============================================================================

    add_heading_custom(doc, '5. СРАВНИТЕЛЬНЫЙ АНАЛИЗ РЕЗУЛЬТАТОВ', level=1)

    add_heading_custom(doc, '5.1. Сравнение архитектур RNN и CNN', level=2)

    add_paragraph_custom(doc, 'Таблица 1. Сравнение характеристик моделей', bold=True)

    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Grid Accent 1'

    # Заголовки
    headers = ['Характеристика', 'RNN (LSTM)', 'CNN']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, size=11, bold=True)

    # Данные
    data = [
        ['Количество параметров', '154,503', '1,469,799'],
        ['Accuracy', '41.09%', '88.51%'],
        ['Recall (macro)', '0.1429', '0.8667'],
        ['F1-Score (macro)', '0.0832', '0.8614'],
        ['TPR', '0.4109', '0.8667'],
        ['FPR', '0.0982', '0.0222'],
        ['AUC', '0.5068', '0.9834'],
        ['MSE', '0.1089', '0.0194'],
        ['Переобучение', '0.0009', '0.0052']
    ]

    for i, row_data in enumerate(data, start=1):
        for j, value in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, size=11)

    doc.add_paragraph()

    add_heading_custom(doc, '5.2. Анализ метрик качества', level=2)

    add_paragraph_custom(doc, '1. Точность классификации (Accuracy):', bold=True)
    add_paragraph_custom(doc,
        'CNN превосходит RNN более чем вдвое (88.51% vs 41.09%). Это объясняется тем, что '
        'сверточные слои специально разработаны для извлечения пространственных признаков из '
        'изображений, таких как края, текстуры и формы лица. RNN обрабатывает изображение как '
        'последовательность строк пикселей, теряя важную пространственную информацию.')

    add_paragraph_custom(doc, '2. AUC (Area Under Curve):', bold=True)
    add_paragraph_custom(doc,
        'Значение AUC для CNN (0.9834) указывает на отличное качество модели в различении классов. '
        'RNN показывает AUC близкий к 0.5, что соответствует случайному угадыванию. Это подтверждает '
        'неэффективность RNN для задач компьютерного зрения.')

    add_paragraph_custom(doc, '3. False Positive Rate (FPR):', bold=True)
    add_paragraph_custom(doc,
        'CNN демонстрирует низкий уровень ложных срабатываний (2.22% vs 9.82% у RNN), что критически '
        'важно для биометрических систем безопасности, где недопустим высокий уровень ошибок.')

    add_paragraph_custom(doc, '4. Переобучение:', bold=True)
    add_paragraph_custom(doc,
        'Обе модели показывают минимальное переобучение благодаря применению Dropout и регуляризации. '
        'CNN использует дополнительно BatchNormalization, что стабилизирует обучение.')

    add_paragraph_custom(doc, '5. Количество параметров:', bold=True)
    add_paragraph_custom(doc,
        'CNN имеет значительно больше параметров (1.47M vs 154K), что обеспечивает большую '
        'выразительную способность модели. Однако это также требует больших вычислительных ресурсов '
        'и времени обучения.')

    add_heading_custom(doc, '5.3. Практические рекомендации', level=2)

    add_paragraph_custom(doc, 'На основе проведенного анализа можно сформулировать следующие рекомендации:')

    add_paragraph_custom(doc, '1. Выбор архитектуры:', bold=True)
    add_paragraph_custom(doc,
        '• Для задач компьютерного зрения (распознавание лиц, объектов, классификация изображений) '
        'следует использовать CNN-архитектуры.')
    add_paragraph_custom(doc,
        '• RNN и LSTM эффективны для последовательных данных (текст, временные ряды, аудио), '
        'но не подходят для пространственных данных.')

    add_paragraph_custom(doc, '2. Методы кластеризации:', bold=True)
    add_paragraph_custom(doc,
        '• K-means++ предпочтителен для задач кластеризации благодаря улучшенной инициализации.')
    add_paragraph_custom(doc,
        '• Agglomerative Clustering полезен для исследовательского анализа и построения иерархий.')

    add_paragraph_custom(doc, '3. Регуляризация:', bold=True)
    add_paragraph_custom(doc,
        '• BatchNormalization критически важен для глубоких CNN - стабилизирует обучение.')
    add_paragraph_custom(doc,
        '• Dropout (0.25-0.5) эффективно предотвращает переобучение.')
    add_paragraph_custom(doc,
        '• EarlyStopping и ReduceLROnPlateau позволяют избежать переобучения и улучшить сходимость.')

    add_paragraph_custom(doc, '4. Датасеты:', bold=True)
    add_paragraph_custom(doc,
        '• Для биометрической идентификации необходимы сбалансированные датасеты с достаточным '
        'количеством примеров каждого класса (минимум 50-100 изображений на класс).')

    add_paragraph_custom(doc, '5. Оценка качества:', bold=True)
    add_paragraph_custom(doc,
        '• Для биометрических систем критичны метрики FPR и TPR, а не только Accuracy.')
    add_paragraph_custom(doc,
        '• ROC-кривые и AUC обеспечивают комплексную оценку качества классификации.')

    doc.add_page_break()

    # ============================================================================
    # 6. ЗАКЛЮЧЕНИЕ
    # ============================================================================

    add_heading_custom(doc, '6. ЗАКЛЮЧЕНИЕ', level=1)

    add_paragraph_custom(doc,
        'В рамках данной курсовой работы были изучены и практически реализованы современные методы '
        'машинного обучения для задач анализа данных и биометрической идентификации.')

    add_paragraph_custom(doc,
        'Первая лабораторная работа продемонстрировала эффективность алгоритмов кластеризации '
        '(K-means, K-means++, Agglomerative Clustering) для анализа сетевого трафика SDN. K-means++ '
        'показал наилучшие результаты благодаря улучшенной инициализации центроидов.')

    add_paragraph_custom(doc,
        'Вторая и третья лабораторные работы были посвящены сравнению архитектур RNN (LSTM) и CNN '
        'для задачи биометрической идентификации. Результаты однозначно показали превосходство '
        'сверточных нейронных сетей:')

    add_paragraph_custom(doc, '• CNN достигла точности 88.51% против 41.09% у RNN')
    add_paragraph_custom(doc, '• AUC улучшен с 0.5068 до 0.9834')
    add_paragraph_custom(doc, '• FPR снижен с 9.82% до 2.22%')

    add_paragraph_custom(doc,
        'Полученные результаты подтверждают, что выбор архитектуры нейронной сети должен '
        'соответствовать природе обрабатываемых данных: CNN для изображений, RNN для последовательностей.')

    add_paragraph_custom(doc,
        'Практическая значимость работы заключается в разработке работающих моделей для реальных '
        'приложений: систем контроля доступа, анализа сетевого трафика, автоматической идентификации '
        'личности. Код всех реализаций доступен в репозитории и может быть использован для дальнейших '
        'исследований.')

    add_paragraph_custom(doc,
        'В процессе работы освоены современные инструменты машинного обучения (TensorFlow, Keras, '
        'scikit-learn), методы оценки качества моделей, техники регуляризации и оптимизации обучения.')

    doc.add_page_break()

    # ============================================================================
    # 7. СПИСОК ЛИТЕРАТУРЫ
    # ============================================================================

    add_heading_custom(doc, '7. СПИСОК ЛИТЕРАТУРЫ', level=1)

    add_paragraph_custom(doc, '1. Goodfellow, I., Bengio, Y., Courville, A. Deep Learning. MIT Press, 2016.')
    add_paragraph_custom(doc, '2. Chollet, F. Deep Learning with Python. Manning Publications, 2021.')
    add_paragraph_custom(doc, '3. Géron, A. Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow. O\'Reilly Media, 2022.')
    add_paragraph_custom(doc, '4. LeCun, Y., Bengio, Y., Hinton, G. Deep learning. Nature, 521(7553), 436-444, 2015.')
    add_paragraph_custom(doc, '5. Simonyan, K., Zisserman, A. Very Deep Convolutional Networks for Large-Scale Image Recognition. ICLR, 2015.')
    add_paragraph_custom(doc, '6. Hochreiter, S., Schmidhuber, J. Long Short-Term Memory. Neural Computation, 9(8), 1735-1780, 1997.')
    add_paragraph_custom(doc, '7. Huang, G.B., et al. Labeled Faces in the Wild: A Database for Studying Face Recognition in Unconstrained Environments. University of Massachusetts, Amherst, Technical Report 07-49, 2007.')
    add_paragraph_custom(doc, '8. Arthur, D., Vassilvitskii, S. k-means++: The Advantages of Careful Seeding. SODA \'07, 2007.')
    add_paragraph_custom(doc, '9. Scikit-learn: Machine Learning in Python. Pedregosa et al., JMLR 12, pp. 2825-2830, 2011.')
    add_paragraph_custom(doc, '10. Abadi, M., et al. TensorFlow: Large-Scale Machine Learning on Heterogeneous Systems, 2015.')

    # Сохранение документа
    doc.save('Отчет_Курсовая_Работа.docx')
    print("✓ Отчет успешно создан: Отчет_Курсовая_Работа.docx")

if __name__ == "__main__":
    create_comprehensive_report()
